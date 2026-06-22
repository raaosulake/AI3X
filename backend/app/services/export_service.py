import os
import markdown
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
from typing import Optional
from app.config import HISTORY_DIR

try:
    from weasyprint import HTML
    WEASYPRINT_AVAILABLE = True
except ImportError:
    WEASYPRINT_AVAILABLE = False


class ExportService:

    def export_markdown(self, content: str, filename: str) -> Optional[str]:
        filepath = HISTORY_DIR / f"{filename}.md"
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(content)
        return str(filepath)

    def export_pdf(self, content: str, filename: str) -> Optional[str]:
        html_content = markdown.markdown(content, extensions=["tables", "fenced_code"])
        styled_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <style>
        body {{ font-family: 'Segoe UI', Arial, sans-serif; margin: 40px; line-height: 1.6; color: #333; }}
        h1 {{ color: #1a237e; border-bottom: 2px solid #1a237e; padding-bottom: 8px; }}
        h2 {{ color: #283593; margin-top: 24px; }}
        h3 {{ color: #3949ab; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px 12px; text-align: left; }}
        th {{ background-color: #e8eaf6; }}
        code {{ background: #f5f5f5; padding: 2px 6px; border-radius: 3px; font-size: 0.9em; }}
        pre {{ background: #f5f5f5; padding: 16px; border-radius: 4px; overflow-x: auto; }}
        ul, ol {{ margin: 8px 0; }}
    </style>
</head>
<body>{html_content}</body></html>"""
        filepath = HISTORY_DIR / f"{filename}.pdf"
        if not WEASYPRINT_AVAILABLE:
            return None
        try:
            HTML(string=styled_html).write_pdf(str(filepath))
            return str(filepath)
        except Exception:
            return None

    def export_docx(self, content: str, filename: str) -> Optional[str]:
        html_content = markdown.markdown(content, extensions=["tables", "fenced_code"])
        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Segoe UI"
        font.size = Pt(11)

        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.children:
            if element.name == "h1":
                p = doc.add_heading(element.get_text(), level=1)
            elif element.name == "h2":
                p = doc.add_heading(element.get_text(), level=2)
            elif element.name == "h3":
                p = doc.add_heading(element.get_text(), level=3)
            elif element.name == "p":
                p = doc.add_paragraph(element.get_text())
            elif element.name == "ul":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "ol":
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Number")
            elif element.name == "table":
                rows = element.find_all("tr")
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0].find_all(["td", "th"])))
                    table.style = "Light Grid Accent 1"
                    for i, row in enumerate(rows):
                        cells = row.find_all(["td", "th"])
                        for j, cell in enumerate(cells):
                            table.rows[i].cells[j].text = cell.get_text()
            elif element.name == "pre":
                p = doc.add_paragraph(element.get_text())
                p.style.font.name = "Consolas"
                p.style.font.size = Pt(9)

        filepath = HISTORY_DIR / f"{filename}.docx"
        doc.save(str(filepath))
        return str(filepath)


export_service = ExportService()
